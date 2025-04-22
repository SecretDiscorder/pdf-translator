import pdfplumber
from docx import Document
from docx.shared import Pt
import sys
import os
import nltk
import pytesseract
from deep_translator import GoogleTranslator
from textblob import TextBlob
import time

nltk_data_dir = "/storage/internal/nltk_data"
# Tentukan lokasi folder resource NLTK
nltk.data.path.append(nltk_data_dir)
nltk.download('punkt', download_dir=nltk_data_dir)
nltk.download('punkt_tab', download_dir=nltk_data_dir)
from deep_translator import GoogleTranslator
from textblob import TextBlob

def safe_translate(text, source_lang, target_lang, delay=1):
    try:
        return GoogleTranslator(source=source_lang, target=target_lang).translate(text)
    except Exception as e:
        print(f"[GoogleTranslator error] {e} — falling back to TextBlob")
        try:
            blob = TextBlob(text)
            return str(blob.translate(from_lang=source_lang if source_lang != "auto" else "en", to=target_lang))
        except Exception as e2:
            print(f"[TextBlob error] {e2} — keeping original text.")
            return text

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    return (
        0x20 <= codepoint <= 0xD7FF
        or codepoint in (0x9, 0xA, 0xD)
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
    )

def clean_text(text):
    return ''.join(c for c in text if valid_xml_char_ordinal(c))

def split_into_sentences(text):
    # Menggunakan nltk untuk membagi teks menjadi kalimat
    return nltk.sent_tokenize(text)
    
    
def pdf_to_docx(pdf_path, output_name="output.docx", source_lang="auto", target_lang="en"):
    if not os.path.exists(pdf_path):
        print("PDF file not found.")
        sys.exit(1)

    doc = Document()
    doc.add_heading(f"Translated PDF ({source_lang} → {target_lang})", 0)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)

        for i, page in enumerate(pdf.pages):
            progress = ((i + 1) / total_pages) * 100
            print(f"Translating page {i + 1}/{total_pages} ({progress:.2f}%)")

            doc.add_heading(f"Page {i + 1}", level=2)
            text = page.extract_text()
            full_paragraph = ""

            if text:
                full_paragraph += " ".join(text.splitlines())

            image = page.to_image(resolution=300)
            pil_image = image.original
            ocr_text = pytesseract.image_to_string(pil_image, lang='vie+jpn')  # if working with Vietnamese
            if ocr_text:
                full_paragraph += " " + ocr_text

            sentences = split_into_sentences(full_paragraph)


            # Process and batch translate
            cleaned_sentences = [clean_text(s).strip() for s in sentences if clean_text(s).strip()]

                
            for sentence in cleaned_sentences:
                try:
                    translated = safe_translate(sentence, source_lang, target_lang)
                    para = doc.add_paragraph(translated)
                    para.style = style
                except Exception as e:
                    print(f"[General translation error] {e}")
                    doc.add_paragraph(sentence)
    
                time.sleep(0.5) 


    doc.save(output_name)
    print(f"Translated and saved to {output_name}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 script.py <file.pdf> [output.docx] [source_lang] [target_lang]")
        sys.exit(1)

    input_pdf = sys.argv[1]
    output_docx = sys.argv[2] if len(sys.argv) > 2 else "output.docx"
    source_lang = sys.argv[3] if len(sys.argv) > 3 else "auto"
    target_lang = sys.argv[4] if len(sys.argv) > 4 else "en"

    pdf_to_docx(input_pdf, output_docx, source_lang, target_lang)
    