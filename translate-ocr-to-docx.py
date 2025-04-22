import pdfplumber
from docx import Document
from docx.shared import Pt
import sys
import os
import nltk

nltk.download('punkt')
nltk.download('punkt_tab')
# Tentukan lokasi folder resource NLTK
nltk.data.path.append("/storage/internal/nltk_data")

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    return (
        0x20 <= codepoint <= 0xD7FF
        or codepoint in (0x9, 0xA, 0xD)
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
    )

def clean_text(text):
    return ''.join(c for c in text if valid_xml_char_ordinal(c))

def split_into_sentences(text):
    # Menggunakan nltk untuk membagi teks menjadi kalimat
    return nltk.sent_tokenize(text)

def pdf_to_docx(pdf_path, output_name="output.docx"):
    if not os.path.exists(pdf_path):
        print("PDF file not found.")
        sys.exit(1)

    doc = Document()
    doc.add_heading("Converted from PDF", 0)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                doc.add_heading(f"Page {i + 1}", level=2)

                # Gabungkan baris jadi satu paragraf
                full_paragraph = " ".join(text.splitlines())
                sentences = split_into_sentences(full_paragraph)

                for sentence in sentences:
                    cleaned_sentence = clean_text(sentence).strip()
                    if cleaned_sentence:
                        para = doc.add_paragraph(cleaned_sentence)
                        para.style = style

                doc.add_page_break()

    doc.save(output_name)
    print(f"Saved to {output_name}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 script.py <file.pdf> [output.docx]")
        sys.exit(1)

    input_pdf = sys.argv[1]
    output_docx = sys.argv[2] if len(sys.argv) > 2 else "output.docx"
    pdf_to_docx(input_pdf, output_docx)
